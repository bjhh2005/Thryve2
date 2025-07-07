from .MessageNode import MessageNode
import os
from typing import Dict, Any, List, Optional
from PyPDF2 import PdfReader, PdfWriter, PdfMerger
from PIL import Image
import fitz  # PyMuPDF
import io
import re
from docx import Document
from docx.shared import Inches, Pt

class PdfProcessor(MessageNode):
    def __init__(self, id: str, type: str, nextNodes: List, eventBus, data: Dict):
        """
        初始化PDF处理节点
        
        Args:
            id (str): 节点ID
            type (str): 节点类型
            nextNodes (list): 下一个节点列表
            eventBus: 事件总线
            data (dict): 节点数据
        """
        super().__init__(id, type, nextNodes, eventBus)
        self.data = data
        self.mode = data.get('mode', '')
        self.input_file = None
        
        # 获取输出路径配置
        self.output_folder = self._get_input_value(data, 'outputFolder') or "output"
        self.output_name = self._get_input_value(data, 'outputName') or "output"
        
        # 确保输出目录存在
        if not os.path.exists(self.output_folder):
            os.makedirs(self.output_folder)
            
        # 对于多文件输出的操作，创建以output_name命名的子文件夹
        if self.mode in ['extract', 'split', 'convert']:
            self.output_dir = os.path.join(self.output_folder, self.output_name)
            if not os.path.exists(self.output_dir):
                os.makedirs(self.output_dir)
        else:
            self.output_dir = self.output_folder

    def _get_input_value(self, data: Dict[str, Any], key: str) -> Any:
        """从inputsValues中获取值，处理constant和ref两种类型"""
        if key not in data.get('inputsValues', {}):
            return None
        
        value_data = data['inputsValues'][key]
        
        # 如果value_data不是字典类型，直接返回其值
        if not isinstance(value_data, dict):
            return value_data
        
        if value_data.get("type") == "constant":
            return value_data.get("content")
        elif value_data.get("type") == "ref":
            content = value_data.get("content", [])
            if len(content) >= 2:
                node_id = content[0]
                if node_id.endswith("_locals"):
                    node_id = node_id[:-7]
                param_name = content[1]
                return self._eventBus.emit("askMessage", node_id, param_name)
        return None

    def run(self) -> bool:
        """
        执行PDF处理节点
        
        Returns:
            bool: 执行是否成功
        """
        try:
            # 从data中获取输入文件
            self.input_file = self._get_input_value(self.data, 'inputFile')
            

            # 根据不同模式执行相应操作
            if self.mode == 'extract':
                result = self._extract_pdf()
            elif self.mode == 'merge':
                result = self._merge_pdfs()
            elif self.mode == 'split':
                result = self._split_pdf()
            elif self.mode == 'convert':
                result = self._convert_pdf()
            elif self.mode == 'compress':
                result = self._compress_pdf()
            elif self.mode == 'encrypt':
                result = self._encrypt_pdf()
            elif self.mode == 'decrypt':
                result = self._decrypt_pdf()
            elif self.mode == 'watermark':
                result = self._add_watermark()
            elif self.mode == 'metadata':
                result = self._edit_metadata()
            else:
                raise ValueError(f"不支持的操作模式: {self.mode}",10)

            # 更新消息
            self.MessageList = result
            
            # 更新下一个节点
            self.updateNext()
            return self.MessageList
            
        except Exception as e:
            raise Exception(f"PDF处理节点 {self._id} 执行错误: {str(e)}", 10)

    def updateNext(self):
        """更新下一个节点"""
        if not self._nextNodes and not self._is_loop_internal:
            raise Exception(f"节点 {self._id}: 缺少后续节点配置", 10)
        self._next = self._nextNodes[0][1]

    def _extract_pdf(self) -> Dict[str, Any]:
        """提取PDF文本和图片"""
        try:
            page_range = self._get_input_value(self.data, 'pageRange')
            extract_images = self._get_input_value(self.data, 'extractImages')
            
            # 打开PDF文件
            doc = fitz.open(self.input_file)
            text_content = []
            images = []
            
            # 处理页面范围
            if page_range:
                pages = self._parse_page_range(page_range, doc.page_count)
            else:
                pages = range(doc.page_count)
            
            # 提取文本和图片
            for page_num in pages:
                page = doc[page_num]
                try:
                    page_text = page.get_text()
                    # 处理可能的编码问题
                    if isinstance(page_text, str):
                        try:
                            # 测试能否正确编码
                            page_text.encode('utf-8')
                        except UnicodeEncodeError:
                            # 如果编码失败，使用replace策略
                            page_text = page_text.encode('utf-8', errors='replace').decode('utf-8')
                    text_content.append(page_text)
                except Exception as e:
                    self._eventBus.emit("message", "warning", self._id, f"提取页面 {page_num + 1} 文本时出错: {str(e)}")
                    text_content.append(f"[页面 {page_num + 1} 文本提取失败]")
                
                if extract_images:
                    for img_index, img in enumerate(page.get_images()):
                        try:
                            xref = img[0]
                            base_image = doc.extract_image(xref)
                            image_data = base_image["image"]
                            
                            # 保存图片到指定目录
                            image_filename = f"page_{page_num + 1}_img_{img_index + 1}.png"
                            image_path = os.path.join(self.output_dir, image_filename)
                            image_path = self._get_unique_filename(image_path)
                            with open(image_path, "wb") as img_file:
                                img_file.write(image_data)
                            images.append(image_path)
                        except Exception as img_err:
                            self._eventBus.emit("message", "warning", self._id, 
                                              f"提取页面 {page_num + 1} 图片 {img_index + 1} 时出错: {str(img_err)}")
            
            # 保存提取的文本到指定目录
            text_content_str = "\n".join(text_content)
            text_output_file = os.path.join(self.output_dir, "extracted_text.txt")
            text_output_file = self._get_unique_filename(text_output_file)
            
            # 使用utf-8编码保存文本，处理可能的编码问题
            try:
                with open(text_output_file, "w", encoding="utf-8", errors="replace") as f:
                    f.write(text_content_str)
            except Exception as write_err:
                self._eventBus.emit("message", "error", self._id, f"保存提取文本时出错: {str(write_err)}")
                # 尝试使用不同的编码方式
                with open(text_output_file, "w", encoding="cp936", errors="replace") as f:
                    f.write(text_content_str)

            doc.close()
            self._eventBus.emit("message", "info", self._id, "PDF extraction completed successfully!")
            return {
                "text": text_content_str,
                "textFile": text_output_file,
                "images": images if extract_images else []
            }
        except Exception as e:
            raise RuntimeError(f"提取PDF内容时发生错误: {str(e)}",10)

    def _merge_pdfs(self) -> Dict[str, Any]:
        """合并多个PDF文件"""
        try:
            input_files = self._get_input_value(self.data, 'inputFiles')
            sort_by = self._get_input_value(self.data, 'sortBy') or 'name'
            
            if sort_by == 'name':
                input_files.sort()
            else:  # sort_by == 'date'
                input_files.sort(key=lambda x: os.path.getmtime(x))
            
            merger = PdfMerger()
            for pdf_file in input_files:
                merger.append(pdf_file)
            
            # 使用指定的输出路径
            output_file = os.path.join(self.output_dir, f"{self.output_name}.pdf")
            output_file = self._get_unique_filename(output_file)
            merger.write(output_file)
            merger.close()
            
            # 计算总页数
            reader = PdfReader(output_file)
            total_pages = len(reader.pages)
            
            self._eventBus.emit("message", "info", self._id, "PDF merge completed successfully!")
            return {
                "outputFile": output_file,
                "pageCount": total_pages
            }
        except Exception as e:
            raise RuntimeError(f"合并PDF文件时发生错误: {str(e)}",10)

    def _split_pdf(self) -> Dict[str, Any]:
        """拆分PDF文件"""
        try:
            split_method = self._get_input_value(self.data, 'splitMethod')
            value = self._get_input_value(self.data, 'value')
            
            reader = PdfReader(self.input_file)
            total_pages = len(reader.pages)
            output_files = []
            
            if split_method == 'byPage':
                pages_per_file = int(value)
                num_files = (total_pages + pages_per_file - 1) // pages_per_file
                
                for i in range(num_files):
                    writer = PdfWriter()
                    start_page = i * pages_per_file
                    end_page = min((i + 1) * pages_per_file, total_pages)
                    
                    for page_num in range(start_page, end_page):
                        writer.add_page(reader.pages[page_num])
                    
                    output_file = os.path.join(self.output_dir, f"{self.output_name}_{i + 1}.pdf")
                    output_file = self._get_unique_filename(output_file)
                    with open(output_file, "wb") as output:
                        writer.write(output)
                    output_files.append(output_file)
            
            elif split_method == 'byBookmark':
                # 使用PyMuPDF处理书签
                doc = fitz.open(self.input_file)
                toc = doc.get_toc()
                if not toc:
                    raise ValueError("PDF文件没有书签")
                
                current_writer = PdfWriter()
                current_start = 0
                file_counter = 1
                
                for bookmark in toc:
                    level, title, page = bookmark
                    if level == 1:  # 只在顶级书签处分割
                        if current_start < page - 1:
                            # 保存当前部分
                            output_file = os.path.join(self.output_dir, f"{self.output_name}_{file_counter}.pdf")
                            with open(output_file, "wb") as output:
                                current_writer.write(output)
                            output_files.append(output_file)
                            
                            current_writer = PdfWriter()
                            current_start = page - 1
                            file_counter += 1
                        
                        # 添加页面到当前writer
                        for p in range(current_start, page):
                            current_writer.add_page(reader.pages[p])
                
                # 保存最后一部分
                if current_writer.pages:
                    output_file = os.path.join(self.output_dir, f"{self.output_name}_{file_counter}.pdf")
                    with open(output_file, "wb") as output:
                        current_writer.write(output)
                    output_files.append(output_file)
                
                doc.close()
            
            elif split_method == 'bySize':
                target_size_mb = float(value)
                current_writer = PdfWriter()
                current_size = 0
                file_counter = 1
                
                for page in reader.pages:
                    # 创建临时文件来估算页面大小
                    temp_writer = PdfWriter()
                    temp_writer.add_page(page)
                    temp_file = os.path.join(self.output_dir, "temp.pdf")
                    with open(temp_file, "wb") as output:
                        temp_writer.write(output)
                    page_size = os.path.getsize(temp_file) / (1024 * 1024)  # 转换为MB
                    os.remove(temp_file)
                    
                    if current_size + page_size > target_size_mb:
                        # 保存当前文件
                        output_file = os.path.join(self.output_dir, f"{self.output_name}_{file_counter}.pdf")
                        with open(output_file, "wb") as output:
                            current_writer.write(output)
                        output_files.append(output_file)
                        
                        current_writer = PdfWriter()
                        current_size = 0
                        file_counter += 1
                    
                    current_writer.add_page(page)
                    current_size += page_size
                
                # 保存最后一个文件
                if current_writer.pages:
                    output_file = os.path.join(self.output_dir, f"{self.output_name}_{file_counter}.pdf")
                    with open(output_file, "wb") as output:
                        current_writer.write(output)
                    output_files.append(output_file)
            
            self._eventBus.emit("message", "info", self._id, "PDF split completed successfully!")
            return {
                "outputFiles": output_files,
                "fileCount": len(output_files)
            }
        except Exception as e:
            raise RuntimeError(f"拆分PDF文件时发生错误: {str(e)}",10)

    def _convert_pdf(self) -> Dict[str, Any]:
        """转换PDF为其他格式"""
        try:
            output_format = self._get_input_value(self.data, 'outputFormat')
            dpi = self._get_input_value(self.data, 'dpi') or 300
            
            doc = fitz.open(self.input_file)
            output_files = []
            conversion_log = []
            
            if output_format == 'docx':
                # 创建新的Word文档
                word_doc = Document()
                
                # 遍历PDF的每一页
                for page_num in range(len(doc)):
                    page = doc[page_num]
                    
                    # 提取文本
                    text = page.get_text()
                    word_doc.add_paragraph(text)
                    
                    # 提取图片
                    for img_index, img in enumerate(page.get_images()):
                        try:
                            xref = img[0]
                            base_image = doc.extract_image(xref)
                            image_data = base_image["image"]
                            
                            # 保存临时图片文件
                            temp_img_path = os.path.join(self.output_dir, f"temp_img_{page_num}_{img_index}.png")
                            with open(temp_img_path, "wb") as img_file:
                                img_file.write(image_data)
                            
                            # 将图片添加到文档
                            word_doc.add_picture(temp_img_path, width=Inches(6.0))
                            
                            # 删除临时图片文件
                            os.remove(temp_img_path)
                        except Exception as img_error:
                            conversion_log.append(f"Warning: Failed to process image on page {page_num + 1}: {str(img_error)}")
                    
                    # 在每页之间添加分页符
                    if page_num < len(doc) - 1:
                        word_doc.add_page_break()
                
                # 保存Word文档
                docx_filename = f"{self.output_name}.docx"
                docx_path = os.path.join(self.output_dir, docx_filename)
                docx_path = self._get_unique_filename(docx_path)
                word_doc.save(docx_path)
                
                output_files.append(docx_path)
                conversion_log.append(f"PDF converted to DOCX: {docx_path}")
            
            else:
                for page_num in range(len(doc)):
                    page = doc[page_num]
                    if output_format in ['png', 'jpg']:
                        pix = page.get_pixmap(matrix=fitz.Matrix(dpi/72, dpi/72))
                        image_filename = f"{self.output_name}_page_{page_num + 1}.{output_format}"
                        image_path = os.path.join(self.output_dir, image_filename)
                        image_path = self._get_unique_filename(image_path)
                        pix.save(image_path)
                        output_files.append(image_path)
                        conversion_log.append(f"Page {page_num + 1} converted to {output_format}")
                    
                    elif output_format == 'text':
                        text = page.get_text()
                        text_filename = f"{self.output_name}_page_{page_num + 1}.txt"
                        text_path = os.path.join(self.output_dir, text_filename)
                        text_path = self._get_unique_filename(text_path)
                        with open(text_path, 'w', encoding='utf-8') as f:
                            f.write(text)
                        output_files.append(text_path)
                        conversion_log.append(f"Page {page_num + 1} converted to text")
                    
                    elif output_format == 'html':
                        html = page.get_text("html")
                        html_filename = f"{self.output_name}_page_{page_num + 1}.html"
                        html_path = os.path.join(self.output_dir, html_filename)
                        html_path = self._get_unique_filename(html_path)
                        with open(html_path, 'w', encoding='utf-8') as f:
                            f.write(html)
                        output_files.append(html_path)
                        conversion_log.append(f"Page {page_num + 1} converted to HTML")
            
            doc.close()
            self._eventBus.emit("message", "info", self._id, "PDF conversion completed successfully!")
            return {
                "outputFiles": output_files,
                "conversionLog": "\n".join(conversion_log)
            }
        except Exception as e:
            raise RuntimeError(f"转换PDF文件时发生错误: {str(e)}",10)

    def _compress_pdf(self) -> Dict[str, Any]:
        """压缩PDF文件"""
        try:
            quality = self._get_input_value(self.data, 'quality') or 'medium'
            
            # 质量设置
            quality_settings = {
                'high': {'image_scale': 0.8, 'jpeg_quality': 80},
                'medium': {'image_scale': 0.6, 'jpeg_quality': 60},
                'low': {'image_scale': 0.4, 'jpeg_quality': 40}
            }
            settings = quality_settings[quality]
            
            doc = fitz.open(self.input_file)
            output_file = os.path.join(self.output_dir, f"{self.output_name}.pdf")
            output_file = self._get_unique_filename(output_file)
            
            for page in doc:
                # 压缩页面上的图片
                for img_index, img in enumerate(page.get_images()):
                    try:
                        xref = img[0]
                        base_image = doc.extract_image(xref)
                        if not base_image:
                            continue
                            
                        image_data = base_image["image"]
                        
                        # 使用PIL压缩图片
                        image = Image.open(io.BytesIO(image_data))
                        
                        # 检查图片尺寸
                        if image.size[0] <= 0 or image.size[1] <= 0:
                            continue
                            
                        # 计算新尺寸
                        new_width = max(1, int(image.size[0] * settings['image_scale']))
                        new_height = max(1, int(image.size[1] * settings['image_scale']))
                        
                        # 只有当新尺寸有效时才进行压缩
                        if new_width > 0 and new_height > 0:
                            image = image.resize((new_width, new_height), Image.LANCZOS)
                            
                            # 转换为RGB模式（如果需要）
                            if image.mode in ('RGBA', 'LA'):
                                background = Image.new('RGB', image.size, (255, 255, 255))
                                if image.mode == 'RGBA':
                                    background.paste(image, mask=image.split()[3])
                                else:
                                    background.paste(image, mask=image.split()[1])
                                image = background
                            elif image.mode != 'RGB':
                                image = image.convert('RGB')
                            
                            compressed_image = io.BytesIO()
                            image.save(compressed_image, format='JPEG', 
                                     quality=settings['jpeg_quality'],
                                     optimize=True)
                            
                            # 替换原图片
                            page.delete_image(xref)
                            page.insert_image(page.rect, stream=compressed_image.getvalue())
                    except Exception as img_error:
                        # 记录图片处理错误但继续处理其他图片
                        self._eventBus.emit("message", "warning", self._id, f"处理图片时出现警告: {str(img_error)}")
                        continue
            
            # 使用强压缩选项保存文档
            doc.save(output_file,
                    garbage=4,  # 最大垃圾收集
                    deflate=True,  # 使用deflate压缩
                    clean=True)  # 清理内容
            
            doc.close()
            
            # 计算压缩比
            original_size = os.path.getsize(self.input_file)
            compressed_size = os.path.getsize(output_file)
            compression_ratio = (original_size - compressed_size) / original_size * 100
            
            self._eventBus.emit("message", "info", self._id, "PDF compression completed successfully!")
            return {
                "outputFile": output_file,
                "originalSize": original_size,
                "compressedSize": compressed_size,
                "compressionRatio": round(compression_ratio, 2)
            }
        except Exception as e:
            raise RuntimeError(f"压缩PDF文件时发生错误: {str(e)}",10)

    def _encrypt_pdf(self) -> Dict[str, Any]:
        """加密PDF文件"""
        try:
            password = self._get_input_value(self.data, 'password')
            permissions = self._get_input_value(self.data, 'permissions') or []
            
            reader = PdfReader(self.input_file)
            writer = PdfWriter()
            
            # 复制所有页面
            for page in reader.pages:
                writer.add_page(page)
            
            # 设置权限
            permission_map = {
                'print': 'P',
                'copy': 'C',
                'modify': 'M',
                'annotate': 'A'
            }
            
            permissions_bits = sum(2**i for i, perm in enumerate(permissions) 
                                 if perm in permission_map)
            
            # 加密PDF
            writer.encrypt(password, password, use_128bit=True, 
                         permissions_flag=permissions_bits)
            
            output_file = os.path.join(self.output_dir, f"{self.output_name}.pdf")
            output_file = self._get_unique_filename(output_file)
            with open(output_file, "wb") as output:
                writer.write(output)
            
            self._eventBus.emit("message", "info", self._id, "PDF encryption completed successfully!")
            return {
                "outputFile": output_file,
                "success": True
            }
        except Exception as e:
            raise RuntimeError(f"加密PDF文件时发生错误: {str(e)}",10)

    def _decrypt_pdf(self) -> Dict[str, Any]:
        """解密PDF文件"""
        try:
            password = self._get_input_value(self.data, 'password')
            
            reader = PdfReader(self.input_file)
            if reader.is_encrypted:
                success = reader.decrypt(password)
                if not success:
                    raise ValueError("密码错误")
            
            writer = PdfWriter()
            for page in reader.pages:
                writer.add_page(page)
            
            output_file = os.path.join(self.output_dir, f"{self.output_name}.pdf")
            output_file = self._get_unique_filename(output_file)
            with open(output_file, "wb") as output:
                writer.write(output)
            
            self._eventBus.emit("message", "info", self._id, "PDF decryption completed successfully!")
            return {
                "outputFile": output_file,
                "success": True
            }
        except Exception as e:
            raise RuntimeError(f"解密PDF文件时发生错误: {str(e)}",10)

    def _add_watermark(self) -> Dict[str, Any]:
        """添加水印到PDF"""
        try:
            watermark_text = self._get_input_value(self.data, 'watermarkText')
            opacity = self._get_input_value(self.data, 'opacity') or 30
            position = self._get_input_value(self.data, 'position') or 'center'
            
            doc = fitz.open(self.input_file)
            
            for page in doc:
                # 获取页面尺寸
                rect = page.rect
                
                # 创建临时PDF页面作为水印模板
                temp_doc = fitz.open()
                temp_page = temp_doc.new_page(width=rect.width, height=rect.height)
                
                # 计算水印文本大小（根据页面大小自适应）
                diagonal = (rect.width ** 2 + rect.height ** 2) ** 0.5
                font_size = diagonal * 0.05  # 水印大小为对角线的05%
                
                # 计算不同位置的坐标
                margin = font_size  # 边距为字体大小
                positions = {
                    "center": (rect.width / 4, rect.height / 4),
                    "topLeft":(margin, margin),
                    "topRight": (rect.width/2 - margin, margin),
                    "bottomLeft": (margin, rect.height/2 - margin),
                    "bottomRight":  (rect.width/2 - margin, rect.height/2 - margin)
                }
                
                # 获取位置（如果未指定则默认为center）
                pos = position if position in positions else "center"
                x, y = positions[pos]
                
                # 设置水印颜色和透明度
                alpha = opacity / 100.0  # 将百分比转换为0-1范围
                gray_level = 0.8
                color = (gray_level, gray_level, gray_level, alpha)  # RGBA颜色
                
                # 根据位置决定是否旋转
                if pos == "center":
                    # 中心位置时旋转45度
                    temp_page.insert_text(
                        (x - font_size/2, y),  # 调整位置以确保真正居中
                        watermark_text,
                        fontname="china-s",
                        fontsize=font_size,
                        color=color
                    )
                    temp_page.set_rotation(45)
                else:
                    # 其他位置不旋转
                    temp_page.insert_text(
                        (x, y),
                        watermark_text,
                        fontname="china-s",
                        fontsize=font_size,
                        color=color
                    )
                
                # 将临时页面作为水印叠加到原页面上
                page.show_pdf_page(
                    rect,  # 目标矩形（整个页面）
                    temp_doc,  # 源文档
                    0,  # 源页面号（第一页）
                )
                
                # 关闭临时文档
                temp_doc.close()
                
                # 应用更改
                page.clean_contents()
            
            output_file = os.path.join(self.output_dir, f"{self.output_name}.pdf")
            output_file = self._get_unique_filename(output_file)
            doc.save(output_file)
            doc.close()
            
            self._eventBus.emit("message", "info", self._id, "PDF watermark added successfully!")
            return {
                "outputFile": output_file
            }
        except Exception as e:
            raise RuntimeError(f"添加水印时发生错误: {str(e)}",10)

    def _edit_metadata(self) -> Dict[str, Any]:
        """编辑PDF元数据"""
        try:
            title = self._get_input_value(self.data, 'title')
            author = self._get_input_value(self.data, 'author')
            subject = self._get_input_value(self.data, 'subject')
            keywords = self._get_input_value(self.data, 'keywords')
            
            reader = PdfReader(self.input_file)
            writer = PdfWriter()
            
            # 复制所有页面
            for page in reader.pages:
                writer.add_page(page)
            
            # 更新元数据
            metadata = {
                "/Title": title,
                "/Author": author,
                "/Subject": subject,
                "/Keywords": keywords
            }
            writer.add_metadata(metadata)
            
            output_file = os.path.join(self.output_dir, f"{self.output_name}.pdf")
            output_file = self._get_unique_filename(output_file)
            with open(output_file, "wb") as output:
                writer.write(output)
            
            self._eventBus.emit("message", "info", self._id, "PDF metadata updated successfully!")
            return {
                "success": True,
                "metadata": metadata
            }
        except Exception as e:
            raise RuntimeError(f"更新元数据时发生错误: {str(e)}",10)

    def _parse_page_range(self, page_range: str, total_pages: int) -> List[int]:
        """解析页面范围字符串"""
        pages = set()
        for part in page_range.split(','):
            if '-' in part:
                start, end = map(int, part.split('-'))
                if start < 1:
                    start = 1
                if end > total_pages:
                    end = total_pages
                pages.update(range(start - 1, end))
            else:
                page = int(part)
                if 1 <= page <= total_pages:
                    pages.add(page - 1)
        return sorted(list(pages))

    def _get_unique_filename(self, filepath: str) -> str:
        """
        生成不重复的文件名。如果文件已存在，在文件名后添加序号。
        
        Args:
            filepath (str): 原始文件路径
            
        Returns:
            str: 不重复的文件路径
        """
        if not os.path.exists(filepath):
            return filepath
            
        directory = os.path.dirname(filepath)
        filename = os.path.basename(filepath)
        name, ext = os.path.splitext(filename)
        
        counter = 1
        while True:
            new_filepath = os.path.join(directory, f"{name}_{counter}{ext}")
            if not os.path.exists(new_filepath):
                return new_filepath
            counter += 1
